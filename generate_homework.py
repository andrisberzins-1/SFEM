"""
generate_homework.py — Standalone script to generate EC3 buckling homework.

Produces:
  - 10 HQ section PNG images (student + teacher versions)
  - Student Word document with variant table + section images
  - Teacher Word document with variant table + full solutions

Usage:
    python generate_homework.py
"""

from __future__ import annotations

import io
import math
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Import solvers via sys.path (no __init__.py in module dirs)
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(PROJECT_ROOT / "section_app"))
sys.path.insert(0, str(PROJECT_ROOT / "buckling_app"))

from section_solver import RectanglePart, calculate, SectionResult  # noqa: E402
from buckling_solver import MemberInput, check_member, STEEL_GRADES  # noqa: E402

import matplotlib
matplotlib.use("Agg")  # non-interactive backend
import matplotlib.pyplot as plt  # noqa: E402
import matplotlib.patches as mpatches  # noqa: E402

from docx import Document  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402

# ---------------------------------------------------------------------------
# Output directory
# ---------------------------------------------------------------------------
OUTPUT_DIR = PROJECT_ROOT / "exchange" / "homework"

# ---------------------------------------------------------------------------
# Color palette (RGBA tuples for matplotlib — matches section_app)
# ---------------------------------------------------------------------------
PART_FILL_COLORS = [
    (31/255, 119/255, 180/255, 0.4),
    (255/255, 127/255, 14/255, 0.4),
    (44/255, 160/255, 44/255, 0.4),
    (214/255, 39/255, 40/255, 0.4),
]
PART_BORDER_COLORS = [
    (31/255, 119/255, 180/255, 1.0),
    (255/255, 127/255, 14/255, 1.0),
    (44/255, 160/255, 44/255, 1.0),
    (214/255, 39/255, 40/255, 1.0),
]

# ---------------------------------------------------------------------------
# 10 HQ section definitions
# (bf_b, bf_h, tw, hw, tf_h, inset)
# tf_b is derived: top flange corners at web centerlines
#
# Rules:
#   - Flange thicknesses (bf_h, tf_h): min 20mm, 10mm increments
#   - Flange widths (bf_b, tf_b): min 10mm increments, even numbers
#   - Web thickness (tw): from {10, 20, 30, 40, 60}, can reuse
#   - Web height (hw): min 10mm increments, even, all unique
#   - Each section has a unique overall shape
# ---------------------------------------------------------------------------
SECTION_PARAMS = [
    # (bf_b, bf_h, tw,  hw,  tf_h, inset)
    (300,  20, 10, 260,  30,  60),   # Section 1  tf_b=170
    (340,  30, 20, 280,  40,  70),   # Section 2  tf_b=170
    (360,  40, 30, 320,  20,  80),   # Section 3  tf_b=170
    (380,  20, 40, 340,  50,  90),   # Section 4  tf_b=160
    (400,  30, 60, 300,  30,  100),  # Section 5  tf_b=140
    (420,  50, 10, 360,  20,  110),  # Section 6  tf_b=190
    (440,  20, 20, 380,  40,  120),  # Section 7  tf_b=180
    (460,  40, 30, 400,  30,  130),  # Section 8  tf_b=170
    (320,  30, 40, 240,  50,  50),   # Section 9  tf_b=180
    (480,  50, 60, 420,  20,  140),  # Section 10 tf_b=140
]

# Steel grade cycling for variants
GRADE_CYCLE = ["S235", "S275", "S355", "S420",
               "S235", "S275", "S355", "S420",
               "S235", "S275"]

# Target utilization per variant (varied for realistic homework)
TARGET_UTILS = [0.80, 0.75, 0.85, 0.78, 0.82, 0.73, 0.88, 0.76, 0.84, 0.79]


# ---------------------------------------------------------------------------
# Section construction
# ---------------------------------------------------------------------------
def make_hq_parts(bf_b, bf_h, tw, hw, tf_h, inset):
    """Build 4 RectangleParts for an HQ beam section.

    Top flange corners sit at the centerlines of the webs:
      tf_z_left  = inset + tw/2
      tf_z_right = (bf_b - inset - tw) + tw/2 = bf_b - inset - tw/2
      tf_b       = bf_b - 2*inset - tw
    """
    lw_z = float(inset)
    rw_z = float(bf_b - inset - tw)
    # Top flange: left edge at left-web centerline, right edge at right-web centerline
    tf_z = lw_z + tw / 2.0
    tf_b = rw_z + tw / 2.0 - tf_z   # = bf_b - 2*inset - tw
    return [
        RectanglePart("Bottom flange", float(bf_b), float(bf_h), 0.0, 0.0),
        RectanglePart("Left web", float(tw), float(hw), float(bf_h), lw_z),
        RectanglePart("Right web", float(tw), float(hw), float(bf_h), rw_z),
        RectanglePart("Top flange", float(tf_b), float(tf_h),
                       float(bf_h + hw), float(tf_z)),
    ]


# ---------------------------------------------------------------------------
# Matplotlib section figure — replicates section_app grid-label view
# ---------------------------------------------------------------------------
def _assign_label_tiers(sorted_vals, min_gap):
    """Assign tier 0, 1, 2 for labels. Higher tiers use leaders."""
    tiers = [0] * len(sorted_vals)
    last_on_tier = {0: -1e9, 1: -1e9, 2: -1e9}
    for i, v in enumerate(sorted_vals):
        placed = False
        for t in [0, 1, 2]:
            if (v - last_on_tier[t]) >= min_gap:
                tiers[i] = t
                last_on_tier[t] = v
                placed = True
                break
        if not placed:
            best = max([0, 1, 2], key=lambda t: v - last_on_tier[t])
            tiers[i] = best
            last_on_tier[best] = v
    return tiers


def _get_leader_flags(tiers):
    """If any label in a cluster has tier > 0, ALL adjacent labels get leaders."""
    flags = [t > 0 for t in tiers]
    for i in range(len(tiers)):
        if tiers[i] == 0:
            if (i > 0 and tiers[i - 1] > 0) or \
               (i < len(tiers) - 1 and tiers[i + 1] > 0):
                flags[i] = True
    return flags


def create_section_figure(parts, result, teacher_mode=False, save_path=None):
    """Create a matplotlib figure replicating the section_app cross-section view.

    Uses grid lines at every edge with coordinate labels (tiered to avoid overlap).
    """
    # Collect all edge coordinates
    all_vert_edges = set()   # y-coordinates
    all_horiz_edges = set()  # z-coordinates

    for part in parts:
        y0 = part.y_bot
        y1 = y0 + part.h
        z0 = part.z_left
        z1 = z0 + part.b
        all_vert_edges.update([y0, y1])
        all_horiz_edges.update([z0, z1])

    all_z = sorted(all_horiz_edges)
    all_y = sorted(all_vert_edges)
    z_min, z_max = min(all_z), max(all_z)
    y_min, y_max = min(all_y), max(all_y)
    z_span = z_max - z_min if z_max > z_min else 1
    y_span = y_max - y_min if y_max > y_min else 1
    pad = max(z_span, y_span) * 0.15

    # Grid metrics (adapted from section_app, with larger gaps for static PNG)
    grid_ext = pad * 0.3
    # Use larger gap thresholds than section_app to prevent overlap in static images
    MIN_LABEL_GAP_Z = max(z_span * 0.07, 25) if z_span > 0 else 10
    MIN_LABEL_GAP_Y = max(y_span * 0.07, 25) if y_span > 0 else 10
    plot_data_span = max(y_span, z_span) + 2 * pad * 1.1
    plot_px = 610
    du_per_px = plot_data_span / plot_px if plot_px > 0 else 1
    text_h = 14 * du_per_px
    text_w = 10 * du_per_px * 3
    label_gap = text_h * 0.25
    tier_step = text_h * 1.6   # wider spacing between tiers for static images

    # Figure sizing
    fig_w = 8
    fig_h = fig_w * (y_span + 2.5 * pad) / (z_span + 2.5 * pad)
    fig, ax = plt.subplots(1, 1, figsize=(fig_w, max(fig_h, 5)))
    ax.set_aspect("equal")

    # Draw rectangles
    for i, part in enumerate(parts):
        ci = i % len(PART_FILL_COLORS)
        rect = mpatches.FancyBboxPatch(
            (part.z_left, part.y_bot), part.b, part.h,
            boxstyle="square,pad=0",
            facecolor=PART_FILL_COLORS[ci],
            edgecolor=PART_BORDER_COLORS[ci],
            linewidth=2,
        )
        ax.add_patch(rect)

    # --- Grid lines at every edge ---
    grid_style = dict(color=(0.6, 0.6, 0.6, 0.4), linewidth=0.8, linestyle=":")
    y_grid_end = z_min - grid_ext
    z_grid_end = y_min - grid_ext

    # Horizontal gridlines + Y-axis labels on the left
    sorted_y = sorted(all_vert_edges)
    y_tiers = _assign_label_tiers(sorted_y, MIN_LABEL_GAP_Y)
    y_leaders = _get_leader_flags(y_tiers)

    for idx, y_val in enumerate(sorted_y):
        ax.plot([y_grid_end, z_max + grid_ext], [y_val, y_val], **grid_style)
        tier = y_tiers[idx]
        use_leader = y_leaders[idx]
        x_label = y_grid_end - label_gap - tier * tier_step
        if use_leader:
            y_spread = {0: text_h * 0.5, 1: 0, 2: -text_h * 0.5}.get(tier, 0)
        else:
            y_spread = 0
        y_text = y_val + y_spread
        label_str = f"{y_val:.0f}"

        if use_leader:
            ax.annotate(
                label_str,
                xy=(y_grid_end, y_val),
                xytext=(x_label, y_text),
                fontsize=11, fontweight="bold", color="black",
                ha="right", va="center",
                arrowprops=dict(arrowstyle="-", color="black", lw=0.8),
            )
        else:
            ax.text(x_label, y_text, label_str,
                    fontsize=11, fontweight="bold", color="black",
                    ha="right", va="center")

    # Vertical gridlines + Z-axis labels on the bottom
    sorted_z = sorted(all_horiz_edges)
    z_tiers = _assign_label_tiers(sorted_z, MIN_LABEL_GAP_Z)
    z_leaders = _get_leader_flags(z_tiers)

    for idx, z_val in enumerate(sorted_z):
        ax.plot([z_val, z_val], [z_grid_end, y_max + grid_ext], **grid_style)
        tier = z_tiers[idx]
        use_leader = z_leaders[idx]
        y_label = z_grid_end - label_gap - tier * tier_step
        if use_leader:
            x_spread = {0: -text_w * 0.8, 1: 0, 2: text_w * 0.8}.get(tier, 0)
            x_anc = {0: "right", 1: "center", 2: "left"}.get(tier, "center")
        else:
            x_spread = 0
            x_anc = "center"
        x_label = z_val + x_spread
        label_str = f"{z_val:.0f}"

        if use_leader:
            ax.annotate(
                label_str,
                xy=(z_val, z_grid_end),
                xytext=(x_label, y_label),
                fontsize=11, fontweight="bold", color="black",
                ha=x_anc, va="top",
                arrowprops=dict(arrowstyle="-", color="black", lw=0.8),
            )
        else:
            ax.text(x_label, y_label, label_str,
                    fontsize=11, fontweight="bold", color="black",
                    ha=x_anc, va="top")

    # --- Initial axes at origin (Yi, Zi) ---
    init_ax_style = dict(color="gray", linewidth=1)
    ax.plot([z_min - pad, z_max + pad], [0, 0], **init_ax_style)
    ax.plot([0, 0], [y_min - pad, y_max + pad], **init_ax_style)
    ax.text(z_max + pad * 0.85, 0, "$Y_i$", fontsize=12, color="gray",
            va="bottom", ha="center")
    ax.text(0, y_max + pad * 0.85, "$Z_i$", fontsize=12, color="gray",
            ha="left", va="center")

    # Teacher mode: centroid + centroidal axes
    if teacher_mode:
        # Centroid marker
        ax.plot(result.zc, result.yc, "x", color="red", markersize=12,
                markeredgewidth=2.5, zorder=10)
        ax.text(result.zc + 8, result.yc + 8, "C",
                color="red", fontsize=12, fontweight="bold", zorder=10)

        # Centroidal axes (green horizontal, blue vertical)
        ax.plot([z_min - pad * 0.8, z_max + pad * 0.8],
                [result.yc, result.yc],
                color="green", linewidth=1.5, linestyle="--", zorder=5)
        ax.plot([result.zc, result.zc],
                [y_min - pad * 0.8, y_max + pad * 0.8],
                color="blue", linewidth=1.5, linestyle="--", zorder=5)

        # Axis labels
        ax.text(z_max + pad * 0.85, result.yc, "Y",
                color="green", fontsize=13, fontweight="bold", va="bottom")
        ax.text(result.zc, y_max + pad * 0.85, "Z",
                color="blue", fontsize=13, fontweight="bold", ha="left")

        # Centroid coordinates
        ax.text(result.zc, result.yc - 12,
                f"({result.zc:.1f}, {result.yc:.1f})",
                color="red", fontsize=9, ha="center", va="top")

    # Layout — extra padding on label sides (left/bottom) for tiered labels
    ax.set_xlim(z_min - pad * 1.5, z_max + pad * 1.1)
    ax.set_ylim(y_min - pad * 1.5, y_max + pad * 1.1)
    ax.set_axis_off()
    fig.patch.set_facecolor("white")

    if save_path:
        fig.savefig(save_path, dpi=150, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
    plt.close(fig)
    return fig


# ---------------------------------------------------------------------------
# Find buckling variant parameters
# ---------------------------------------------------------------------------
def find_variant(sec_result: SectionResult, grade_name: str,
                  target_util: float = 0.80):
    """Find L and N_Ed for target utilization by I_min buckling."""
    fy = STEEL_GRADES[grade_name]
    A = sec_result.A_total
    Iy = sec_result.Iy
    Iz = sec_result.Iz

    best = None

    for L in range(3, 11):  # 3 to 10 m
        inp_tmp = MemberInput(
            name="tmp",
            N_Ed_kN=-1.0,
            A_mm2=A, Iy_mm4=Iy, Iz_mm4=Iz,
            fy_MPa=fy, L_m=float(L),
            mu_y=1.0, mu_z=1.0,
            curve_y="c", curve_z="c",
        )
        result_tmp = check_member(inp_tmp)
        N_b_Rd_y = result_tmp.buckling_y.N_b_Rd_kN
        N_b_Rd_z = result_tmp.buckling_z.N_b_Rd_kN
        N_b_Rd_gov = min(N_b_Rd_y, N_b_Rd_z)
        N_Rd = result_tmp.strength.N_Rd_kN

        N_Ed_candidate = round(target_util * N_b_Rd_gov / 5) * 5
        if N_Ed_candidate < 5:
            continue

        util_buck = N_Ed_candidate / N_b_Rd_gov
        util_str = N_Ed_candidate / N_Rd

        if 0.65 <= util_buck <= 0.95 and util_buck > util_str:
            if best is None or abs(util_buck - target_util) < abs(best["util"] - target_util):
                best = {
                    "L": L,
                    "N_Ed": N_Ed_candidate,
                    "util": util_buck,
                    "util_str": util_str,
                }

    if best is None:
        # Fallback: relax to 1 kN steps
        for L in range(3, 11):
            inp_tmp = MemberInput(
                name="tmp", N_Ed_kN=-1.0,
                A_mm2=A, Iy_mm4=Iy, Iz_mm4=Iz,
                fy_MPa=fy, L_m=float(L),
                mu_y=1.0, mu_z=1.0, curve_y="c", curve_z="c",
            )
            result_tmp = check_member(inp_tmp)
            N_b_Rd_gov = min(result_tmp.buckling_y.N_b_Rd_kN,
                             result_tmp.buckling_z.N_b_Rd_kN)
            N_Rd = result_tmp.strength.N_Rd_kN

            N_Ed_candidate = round(target_util * N_b_Rd_gov)
            util_buck = N_Ed_candidate / N_b_Rd_gov
            util_str = N_Ed_candidate / N_Rd

            if 0.60 <= util_buck <= 0.98 and util_buck > util_str:
                best = {"L": L, "N_Ed": N_Ed_candidate,
                        "util": util_buck, "util_str": util_str}
                break

    return best


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------
def set_cell_text(cell, text, bold=False, size=10, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold


def add_variant_table(doc, variants):
    """Add the variant assignment table to a document."""
    table = doc.add_table(rows=1 + len(variants), cols=5, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Variant", "L (m)", "N_Ed (kN)", "Steel Grade", "Section"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for j, v in enumerate(variants):
        row = table.rows[j + 1]
        set_cell_text(row.cells[0], str(j + 1), size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], str(v["L"]), size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], str(v["N_Ed"]), size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], v["grade"], size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], str(j + 1), size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return table


def add_section_props_table(doc, results):
    """Add section properties summary table (teacher doc)."""
    headers = ["Sec", "A (mm\u00b2)", "Iy (mm\u2074)", "Iz (mm\u2074)",
               "iy (mm)", "iz (mm)", "Wy (mm\u00b3)", "Wz (mm\u00b3)",
               "yc (mm)", "zc (mm)"]
    table = doc.add_table(rows=1 + len(results), cols=len(headers),
                          style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for j, r in enumerate(results):
        row = table.rows[j + 1]
        vals = [
            str(j + 1),
            f"{r.A_total:.0f}",
            f"{r.Iy:.0f}",
            f"{r.Iz:.0f}",
            f"{r.iy:.2f}",
            f"{r.iz:.2f}",
            f"{r.Wy:.0f}",
            f"{r.Wz:.0f}",
            f"{r.yc:.1f}",
            f"{r.zc:.1f}",
        ]
        for i, val in enumerate(vals):
            set_cell_text(row.cells[i], val, size=8,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return table


def add_solution_block(doc, variant_num, variant, sec_result, member_result):
    """Add detailed solution for one variant (teacher doc)."""
    doc.add_heading(f"Variant {variant_num}", level=2)

    inp = member_result.input
    s = member_result.strength

    # Input summary
    p = doc.add_paragraph()
    p.add_run("Input: ").bold = True
    p.add_run(
        f"Section {variant_num}, "
        f"L = {inp.L_m:.0f} m, "
        f"N_Ed = {abs(inp.N_Ed_kN):.0f} kN (compression), "
        f"Steel {variant['grade']} (fy = {inp.fy_MPa:.0f} MPa), "
        f"Pinned-pinned (\u03bc = 1.0), Curve c (\u03b1 = 0.49)"
    )

    # Section properties
    p = doc.add_paragraph()
    p.add_run("Section properties: ").bold = True
    p.add_run(
        f"A = {sec_result.A_total:.0f} mm\u00b2, "
        f"Iy = {sec_result.Iy:.0f} mm\u2074, "
        f"Iz = {sec_result.Iz:.0f} mm\u2074, "
        f"iy = {sec_result.iy:.2f} mm, "
        f"iz = {sec_result.iz:.2f} mm"
    )

    # Strength check
    doc.add_heading("Strength Check", level=3)
    p = doc.add_paragraph()
    p.add_run(f"N_Rd = A \u00d7 fy / \u03b3_M0 = "
              f"{sec_result.A_total:.0f} \u00d7 {inp.fy_MPa:.0f} / 1.0 = "
              f"{s.N_Rd_kN:.1f} kN")
    p = doc.add_paragraph()
    p.add_run(f"Utilization (strength) = {abs(inp.N_Ed_kN):.0f} / {s.N_Rd_kN:.1f} = "
              f"{s.utilization:.3f} ({s.utilization*100:.1f}%) \u2014 "
              f"{'OK' if s.passed else 'FAIL'}")

    # Buckling checks
    for bk, axis_name in [(member_result.buckling_y, "Y"),
                          (member_result.buckling_z, "Z")]:
        doc.add_heading(f"Buckling Check \u2014 {axis_name}-axis", level=3)

        steps = [
            f"Step 1: L_cr = \u03bc \u00d7 L = {bk.mu:.1f} \u00d7 {bk.L_m:.0f} = {bk.L_cr_m:.1f} m = {bk.L_cr_mm:.0f} mm",
            f"Step 2: I_{axis_name.lower()} = {bk.I_mm4:.0f} mm\u2074, "
            f"i_{axis_name.lower()} = \u221a(I/A) = \u221a({bk.I_mm4:.0f}/{sec_result.A_total:.0f}) = {bk.i_mm:.2f} mm",
            f"Step 3: N_Rk = A \u00d7 fy = {sec_result.A_total:.0f} \u00d7 {inp.fy_MPa:.0f} / 1000 = {bk.N_Rk_kN:.1f} kN",
            f"Step 4: N_cr = \u03c0\u00b2 \u00d7 E \u00d7 I / L_cr\u00b2 = "
            f"\u03c0\u00b2 \u00d7 {inp.E_MPa:.0f} \u00d7 {bk.I_mm4:.0f} / {bk.L_cr_mm:.0f}\u00b2 / 1000 = {bk.N_cr_kN:.1f} kN",
            f"Step 5: \u03bb = \u221a(N_Rk / N_cr) = \u221a({bk.N_Rk_kN:.1f} / {bk.N_cr_kN:.1f}) = {bk.lambda_bar:.4f}",
            f"Step 6: Buckling curve: {bk.curve}",
            f"Step 7: \u03b1 = {bk.alpha:.2f}",
            f"Step 8: \u03a6 = 0.5 \u00d7 [1 + \u03b1(\u03bb - 0.2) + \u03bb\u00b2] = "
            f"0.5 \u00d7 [1 + {bk.alpha:.2f} \u00d7 ({bk.lambda_bar:.4f} - 0.2) + {bk.lambda_bar:.4f}\u00b2] = {bk.Phi:.4f}",
            f"Step 9: \u03c7 = 1 / (\u03a6 + \u221a(\u03a6\u00b2 - \u03bb\u00b2)) = "
            f"1 / ({bk.Phi:.4f} + \u221a({bk.Phi:.4f}\u00b2 - {bk.lambda_bar:.4f}\u00b2)) = {bk.chi:.4f}",
            f"Step 10: N_b,Rd = \u03c7 \u00d7 A \u00d7 fy / \u03b3_M1 = "
            f"{bk.chi:.4f} \u00d7 {sec_result.A_total:.0f} \u00d7 {inp.fy_MPa:.0f} / 1.0 / 1000 = {bk.N_b_Rd_kN:.1f} kN",
            f"Step 11: Utilization = {bk.N_Ed_kN:.0f} / {bk.N_b_Rd_kN:.1f} = "
            f"{bk.utilization:.3f} ({bk.utilization*100:.1f}%) \u2014 {'OK' if bk.passed else 'FAIL'}",
        ]
        for step in steps:
            doc.add_paragraph(step, style="List Bullet")

    # Governing check
    doc.add_heading("Conclusion", level=3)
    gov = member_result.governing_check.replace("_", " ").title()
    p = doc.add_paragraph()
    r = p.add_run(
        f"Governing check: {gov}, "
        f"Utilization = {member_result.governing_utilization:.3f} "
        f"({member_result.governing_utilization*100:.1f}%) \u2014 "
        f"{'PASS' if member_result.overall_passed else 'FAIL'}"
    )
    r.bold = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    # Fix Windows console encoding
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8",
                                  errors="replace")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # -----------------------------------------------------------------------
    # Step 1: Build sections and calculate properties
    # -----------------------------------------------------------------------
    print("=" * 60)
    print("STEP 1: Building 10 HQ sections")
    print("=" * 60)

    all_parts = []
    all_results = []

    for i, params in enumerate(SECTION_PARAMS):
        parts = make_hq_parts(*params)
        result = calculate(parts)
        all_parts.append(parts)
        all_results.append(result)

        bf_b, bf_h, tw, hw, tf_h, inset = params
        total_h = bf_h + hw + tf_h
        tf_b = bf_b - 2 * inset - tw
        print(f"  Section {i+1:2d}: "
              f"H={total_h}mm, bf={bf_b}x{bf_h}, tw={tw}, hw={hw}, "
              f"tf={tf_b}x{tf_h}, "
              f"A={result.A_total:.0f}mm\u00b2, "
              f"Iy={result.Iy:.0f}, Iz={result.Iz:.0f}, "
              f"yc={result.yc:.1f}, zc={result.zc:.1f}")

    # -----------------------------------------------------------------------
    # Step 2: Export PNG images
    # -----------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("STEP 2: Exporting section images")
    print("=" * 60)

    for i, (parts, result) in enumerate(zip(all_parts, all_results)):
        # Student version
        path_s = OUTPUT_DIR / f"section_{i+1}.png"
        create_section_figure(parts, result, teacher_mode=False,
                              save_path=str(path_s))
        print(f"  Saved: {path_s.name}")

        # Teacher version
        path_t = OUTPUT_DIR / f"section_{i+1}_teacher.png"
        create_section_figure(parts, result, teacher_mode=True,
                              save_path=str(path_t))
        print(f"  Saved: {path_t.name}")

    # -----------------------------------------------------------------------
    # Step 3: Find buckling variants
    # -----------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("STEP 3: Finding buckling variants (target 70-90% utilization)")
    print("=" * 60)

    variants = []
    member_results = []

    for i, (sec_result, grade, tgt) in enumerate(zip(all_results, GRADE_CYCLE,
                                                          TARGET_UTILS)):
        v = find_variant(sec_result, grade, target_util=tgt)
        if v is None:
            print(f"  WARNING: No valid variant found for section {i+1} / {grade}")
            v = {"L": 6, "N_Ed": 100, "util": 0.0, "util_str": 0.0}

        v["grade"] = grade
        v["section_idx"] = i + 1
        variants.append(v)

        # Run full check
        inp = MemberInput(
            name=f"Variant {i+1}",
            N_Ed_kN=-float(v["N_Ed"]),
            A_mm2=sec_result.A_total,
            Iy_mm4=sec_result.Iy,
            Iz_mm4=sec_result.Iz,
            fy_MPa=STEEL_GRADES[grade],
            L_m=float(v["L"]),
            mu_y=1.0, mu_z=1.0,
            curve_y="c", curve_z="c",
        )
        mr = check_member(inp)
        member_results.append(mr)

        print(f"  Variant {i+1:2d}: Section {i+1}, {grade}, "
              f"L={v['L']}m, N_Ed={v['N_Ed']}kN, "
              f"governing={mr.governing_check}, "
              f"util={mr.governing_utilization:.1%}")

    # -----------------------------------------------------------------------
    # Step 4: Student Word document
    # -----------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("STEP 4: Generating student document")
    print("=" * 60)

    doc_s = Document()

    # Title
    title = doc_s.add_heading("Homework \u2014 EC3 Column Buckling Check", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Task description
    doc_s.add_paragraph(
        "For your assigned variant, perform the following calculations by hand:"
    )
    doc_s.add_paragraph(
        "1. Calculate the cross-section properties of the given HQ section: "
        "Area (A), Moments of inertia (Iy, Iz), Section moduli (Wy, Wz), "
        "Radii of gyration (iy, iz), and Centroid position.",
        style="List Number"
    )
    doc_s.add_paragraph(
        "2. Perform a strength check (EN 1993-1-1 cl. 6.2.4) and a buckling check "
        "(EN 1993-1-1 cl. 6.3.1) for both axes. Determine the governing check "
        "and overall utilization ratio.",
        style="List Number"
    )

    # Common parameters
    doc_s.add_heading("Common Parameters (all variants)", level=2)
    doc_s.add_paragraph(
        "Boundary conditions: Pinned-pinned (\u03bc_y = \u03bc_z = 1.0)\n"
        "Buckling curve: c (\u03b1 = 0.49)\n"
        "Partial safety factors: \u03b3_M0 = \u03b3_M1 = 1.0\n"
        "Elastic modulus: E = 200,000 MPa"
    )

    # Variant table
    doc_s.add_heading("Variant Assignments", level=2)
    add_variant_table(doc_s, variants)

    # Section images
    doc_s.add_page_break()
    doc_s.add_heading("Cross-Section Drawings", level=2)
    doc_s.add_paragraph(
        "All dimensions are in millimetres (mm). "
        "Coordinate labels show absolute positions of each edge."
    )
    for i in range(10):
        img_path = OUTPUT_DIR / f"section_{i+1}.png"
        doc_s.add_paragraph(f"Section {i+1}", style="Heading 3")
        doc_s.add_picture(str(img_path), width=Inches(5.0))

    doc_s_path = OUTPUT_DIR / "homework_student.docx"
    doc_s.save(str(doc_s_path))
    print(f"  Saved: {doc_s_path}")

    # -----------------------------------------------------------------------
    # Step 5: Teacher Word document
    # -----------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("STEP 5: Generating teacher document")
    print("=" * 60)

    doc_t = Document()

    title = doc_t.add_heading("Homework \u2014 TEACHER KEY", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section properties table
    doc_t.add_heading("Section Properties Summary", level=2)
    add_section_props_table(doc_t, all_results)

    # Section images (teacher versions)
    doc_t.add_page_break()
    doc_t.add_heading("Cross-Section Drawings (with centroid & axes)", level=2)
    for i in range(10):
        img_path = OUTPUT_DIR / f"section_{i+1}_teacher.png"
        doc_t.add_paragraph(f"Section {i+1}", style="Heading 3")
        doc_t.add_picture(str(img_path), width=Inches(5.0))

    # Variant table
    doc_t.add_page_break()
    doc_t.add_heading("Variant Assignments", level=2)
    add_variant_table(doc_t, variants)

    # Detailed solutions
    doc_t.add_page_break()
    doc_t.add_heading("Detailed Solutions", level=1)
    for i, (v, sec_result, mr) in enumerate(zip(variants, all_results,
                                                 member_results)):
        add_solution_block(doc_t, i + 1, v, sec_result, mr)
        if i < 9:
            doc_t.add_page_break()

    doc_t_path = OUTPUT_DIR / "homework_teacher.docx"
    doc_t.save(str(doc_t_path))
    print(f"  Saved: {doc_t_path}")

    # -----------------------------------------------------------------------
    # Summary
    # -----------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("DONE! All files saved to:", OUTPUT_DIR)
    print("=" * 60)
    files = list(OUTPUT_DIR.glob("*"))
    print(f"  Total files: {len(files)}")
    for f in sorted(files):
        print(f"    {f.name}")


if __name__ == "__main__":
    main()
